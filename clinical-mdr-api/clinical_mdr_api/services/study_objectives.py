from logging import getLogger
from typing import Dict

from docx.enum.style import WD_STYLE_TYPE
from yattag.doc import Doc

from clinical_mdr_api.models.unit_definition import UnitDefinitionModel
from clinical_mdr_api.oauth import get_current_user_id
from clinical_mdr_api.services._meta_repository import MetaRepository
from clinical_mdr_api.services.study import StudyService
from clinical_mdr_api.services.study_endpoint_selection import (
    StudyEndpointSelectionService,
)
from clinical_mdr_api.services.unit_definition import UnitDefinitionService
from clinical_mdr_api.services.utils.docx_builder import DocxBuilder


# TODO LOCALIZATION
def _gettext(x):
    return x


log = getLogger(__name__)

# pylint: disable=no-member
STYLES = {
    "table": ("SB Table Condensed", WD_STYLE_TYPE.TABLE),
    "header1": ("Table Header lvl1", WD_STYLE_TYPE.PARAGRAPH),
    "header2": ("Table Header lvl2", WD_STYLE_TYPE.PARAGRAPH),
    "objective-level": ("Table Header lvl2", WD_STYLE_TYPE.PARAGRAPH),
    "endpoint-level": ("Table lvl 1", WD_STYLE_TYPE.PARAGRAPH),
    "objective": ("Table lvl 2", WD_STYLE_TYPE.PARAGRAPH),
    "endpoint": ("Table lvl 3", WD_STYLE_TYPE.PARAGRAPH),
    "timeframe": ("Table lvl 4", WD_STYLE_TYPE.PARAGRAPH),
    "units": ("Table lvl 4", WD_STYLE_TYPE.PARAGRAPH),
    "ul": ("Bullet List", WD_STYLE_TYPE.PARAGRAPH),
    "ol": ("Bullet List Numbered", WD_STYLE_TYPE.PARAGRAPH),
}


class StudyObjectivesService:
    """Assemble and visualize Study Protocol Flowchart data"""

    def __init__(self, user_id: str) -> None:
        self._current_user_id = user_id
        self._epoch_terms = None
        self._study_endpoint_selection_service = StudyEndpointSelectionService(
            author=self._current_user_id
        )
        self._meta_repository = MetaRepository(user=user_id)
        self._unit_definition_service = UnitDefinitionService(
            user_id=self._current_user_id, meta_repository=self._meta_repository
        )
        self._units = None

    def _get_all_selection(self, study_uid):
        # Check if study exists
        StudyService(user=get_current_user_id()).get_by_uid(uid=study_uid)

        # Get Endpoint Selections
        selection = self._study_endpoint_selection_service.get_all_selection(
            study_uid=study_uid, no_brackets=True
        )

        return selection.items

    def _get_all_units(self) -> Dict[str, UnitDefinitionModel]:
        return {
            u.uid: u
            for u in self._unit_definition_service.get_all(library_name=None).items
        }

    @property
    def units(self) -> Dict[str, UnitDefinitionModel]:
        if self._units is None:
            self._units = self._get_all_units()
        return self._units

    # Unused but kept for future layout
    def get_condensed_html(self, study_uid) -> str:
        selection = self._get_all_selection(study_uid)
        root = self._build_condensed_tree(selection)
        return self._build_condensed_html(root)

    def get_standard_html(self, study_uid) -> str:
        selection = self._get_all_selection(study_uid)
        tree = self._build_tree(selection)
        return self._build_standard_html(tree)

    # Not used but kept for future layout
    def get_condensed_docx(self, study_uid) -> DocxBuilder:
        selection = self._get_all_selection(study_uid)
        root = self._build_condensed_tree(selection)
        return self._build_condensed_docx(root)

    def get_standard_docx(self, study_uid) -> DocxBuilder:
        selection = self._get_all_selection(study_uid)
        tree = self._build_tree(selection)
        return self._build_standard_docx(tree)

    @staticmethod
    def _build_tree(selection) -> Dict:
        root = {}

        for study_selection_endpoint in selection:
            node = root
            study_objective = study_selection_endpoint.study_objective
            if not study_objective:
                continue

            objective_level = study_objective.objective_level
            node = node.setdefault(objective_level.term_uid, (objective_level, {}))

            node = node[1].setdefault(
                study_objective.study_objective_uid, (study_objective, {})
            )

            if study_selection_endpoint.endpoint_sublevel:
                endpoint_level = study_selection_endpoint.endpoint_sublevel
            else:
                endpoint_level = study_selection_endpoint.endpoint_level
            if not endpoint_level:
                continue
            node = node[1].setdefault(endpoint_level.term_uid, (endpoint_level, {}))

            node[1].setdefault(
                study_selection_endpoint.study_endpoint_uid, study_selection_endpoint
            )

        return root

    @staticmethod
    def _build_condensed_tree(selection) -> Dict:
        root = {}

        for study_selection_endpoint in selection:
            node = root
            study_objective = study_selection_endpoint.study_objective
            if not study_objective:
                continue

            objective_level = study_objective.objective_level
            node = node.setdefault(objective_level.term_uid, (objective_level, {}, {}))

            node[1].setdefault(study_objective.study_objective_uid, study_objective)

            if study_selection_endpoint.endpoint_sublevel:
                endpoint_level = study_selection_endpoint.endpoint_sublevel
            else:
                endpoint_level = study_selection_endpoint.endpoint_level
            if not endpoint_level:
                continue
            node = node[2].setdefault(endpoint_level.term_uid, (endpoint_level, {}))

            node[1].setdefault(
                study_selection_endpoint.study_endpoint_uid, study_selection_endpoint
            )

        return root

    @staticmethod
    def _build_condensed_html(tree) -> str:
        doc, tag, _, line = Doc().ttl()
        doc.asis("<!DOCTYPE html>")

        with tag("html", lang="en"):
            with tag("head"):
                line("title", _gettext("Study Objectives"))

            with tag("body"):
                with tag("table", id="ObjectivesEndpointsTable"):

                    with tag("thead"):
                        with tag("tr"):
                            # TODO: Do we have a CTTermName for these?
                            line("th", _gettext("Objectives"))
                            line("th", _gettext("Endpoints"))

                    with tag("tbody"):
                        for objective_level, objectives, endpoints in sorted(
                            tree.values(), key=lambda o: o[0].order
                        ):
                            with tag("tr"):

                                with tag("td"):
                                    line(
                                        "p",
                                        f"{objective_level.sponsor_preferred_name}:",
                                        klass="objective-level",
                                    )
                                    for study_objective in sorted(
                                        objectives.values(), key=lambda o: o.order
                                    ):
                                        doc.asis(study_objective.objective.name)

                                with tag("td"):
                                    for endpoint_level, study_endpoints in sorted(
                                        endpoints.values(), key=lambda e: e[0].order
                                    ):
                                        line(
                                            "p",
                                            f"{endpoint_level.sponsor_preferred_name}:",
                                            klass="endpoint-level",
                                        )
                                        with tag("ul"):
                                            for study_endpoint in sorted(
                                                study_endpoints.values(),
                                                key=lambda e: e.order,
                                            ):
                                                with tag("li"):
                                                    doc.asis(
                                                        study_endpoint.endpoint.name
                                                    )

        return doc.getvalue()

    def _build_standard_html(self, tree) -> str:
        doc, tag, _, line = Doc().ttl()
        doc.asis("<!DOCTYPE html>")

        with tag("html", lang="en"):
            with tag("head"):
                line("title", _gettext("Study Objectives"))

            with tag("body"):
                with tag("table", id="ObjectivesEndpointsTable"):

                    with tag("thead"):
                        with tag("tr"):
                            # TODO: Do we have a CTTermName for these?
                            line("th", _gettext("Objectives"))
                            line("th", _gettext("Endpoints"), colspan=3)

                    with tag("tbody"):
                        for objective_level, study_objectives in sorted(
                            tree.values(), key=lambda o: o[0].order
                        ):
                            with tag("tr"):

                                line(
                                    "th",
                                    objective_level.sponsor_preferred_name,
                                    klass="objective-level",
                                )
                                line("th", _gettext("Title"), klass="header2")
                                line("th", _gettext("Time frame"), klass="header2")
                                line("th", _gettext("Unit"), klass="header2")

                            for study_objective, endpoint_levels in sorted(
                                study_objectives.values(), key=lambda o: o[0].order
                            ):
                                for epl_idx, epl_ste in enumerate(
                                    sorted(
                                        endpoint_levels.values(),
                                        key=lambda o: o[0].order,
                                    )
                                ):
                                    endpoint_level, study_endpoints = epl_ste

                                    if epl_idx == 0:
                                        with tag("tr"):
                                            with tag(
                                                "td",
                                                klass="objective",
                                                rowspan=(len(endpoint_levels) + 1),
                                            ):
                                                doc.asis(study_objective.objective.name)
                                            line(
                                                "th",
                                                endpoint_level.sponsor_preferred_name,
                                                klass="endpoint-level",
                                                colspan=3,
                                            )

                                    for study_endpoint in sorted(
                                        study_endpoints.values(), key=lambda o: o.order
                                    ):
                                        with tag("tr"):
                                            with tag("td", klass="endpoint"):
                                                doc.asis(study_endpoint.endpoint.name)

                                            with tag("td", klass="timeframe"):
                                                if study_endpoint.timeframe:
                                                    doc.asis(
                                                        study_endpoint.timeframe.name
                                                    )

                                            units_text = self._endpoint_units_to_text(
                                                study_endpoint.endpoint_units
                                            )
                                            line("td", units_text, kalss="units")

        return doc.getvalue()

    @staticmethod
    def _build_condensed_docx(tree) -> DocxBuilder:
        docx = DocxBuilder(STYLES)
        table = docx.create_table(num_rows=1, num_columns=2)

        row = table.rows[0]

        # Header text
        row.cells[0].text = _gettext("Objectives")
        row.cells[1].text = _gettext("Endpoints")

        # Apply paragraph style on all cells of the header row
        docx.format_row(row, [STYLES["header1"][0]] * len(row.cells))

        # Set header row to repeat after page breaks
        docx.repeat_table_header(row)

        for objective_level, objectives, endpoints in sorted(
            tree.values(), key=lambda o: o[0].order
        ):
            row = table.add_row()

            cell = row.cells[0]
            cell.add_paragraph(
                f"{objective_level.sponsor_preferred_name}:",
                style=STYLES["objective-level"][0],
            )
            # Remove first empty paragraph added automatically to cell
            docx.delete_paragraph(cell.paragraphs[0])

            for study_objective in sorted(objectives.values(), key=lambda o: o.order):
                docx.add_html(
                    cell, study_objective.objective.name, default_style="objective"
                )

            cell = row.cells[1]
            for endpoint_level, study_endpoints in sorted(
                endpoints.values(), key=lambda e: e[0].order
            ):
                cell.add_paragraph(
                    f"{endpoint_level.sponsor_preferred_name}:",
                    style=STYLES["endpoint-level"][0],
                )
                for study_endpoint in sorted(
                    study_endpoints.values(), key=lambda e: e.order
                ):
                    docx.add_html(
                        cell, study_endpoint.endpoint.name, default_style="endpoint"
                    )
            if len(cell.paragraphs) > 1:
                # Remove first empty paragraph added automatically to cell
                docx.delete_paragraph(cell.paragraphs[0])

        return docx

    def _build_standard_docx(self, tree) -> DocxBuilder:
        num_rows, num_cols = 1, 4

        docx = DocxBuilder(STYLES)
        table = docx.create_table(num_rows=num_rows, num_columns=num_cols)

        row = table.rows[0]

        # Header
        row.cells[0].text = _gettext("Objectives")
        row.cells[1].text = _gettext("Endpoints")
        # Merge 2nd cell to end of row
        docx.merge_cells(row.cells[1:num_cols])

        # Apply paragraph style on all cells of the header row
        docx.format_row(row, [STYLES["header1"][0]] * num_cols)

        # Set header row to repeat after page breaks
        docx.repeat_table_header(row)

        for objective_level, study_objectives in sorted(
            tree.values(), key=lambda o: o[0].order
        ):
            row = table.add_row()
            num_rows += 1

            docx.replace_content(
                row.cells[0],
                str(objective_level.sponsor_preferred_name),
                style="objective-level",
            )
            # TODO Do we have CT-terms for these?
            docx.replace_content(row.cells[1], _gettext("Title"), style="header2")
            docx.replace_content(row.cells[2], _gettext("Time frame"), style="header2")
            docx.replace_content(row.cells[3], _gettext("Unit"), style="header2")

            for study_objective, endpoint_levels in sorted(
                study_objectives.values(), key=lambda o: o[0].order
            ):
                row = table.add_row()
                num_rows += 1

                docx.add_html(
                    row.cells[0],
                    study_objective.objective.name,
                    default_style="objective",
                )

                if len(row.cells[0].paragraphs) > 1:
                    # Remove first empty paragraph of cell
                    docx.delete_paragraph(row.cells[0].paragraphs[0])

                for epl_idx, epl_ste in enumerate(
                    sorted(endpoint_levels.values(), key=lambda o: o[0].order)
                ):
                    endpoint_level, study_endpoints = epl_ste

                    if epl_idx:
                        # Start a new row expect for the first endpoint-level
                        row = table.add_row()
                        num_rows += 1
                        # Merge first column to previous row
                        docx.merge_cells(
                            [table.rows[num_rows - 2].cells[0], row.cells[0]]
                        )

                    docx.replace_content(
                        row.cells[1],
                        str(endpoint_level.sponsor_preferred_name),
                        style="endpoint-level",
                    )
                    # Merge 2nd cell to end of row
                    docx.merge_cells(row.cells[1:num_cols])

                    for study_endpoint in sorted(
                        study_endpoints.values(), key=lambda o: o.order
                    ):
                        row = table.add_row()
                        num_rows += 1

                        # Merge first column to previous row
                        docx.merge_cells(
                            [table.rows[num_rows - 2].cells[0], row.cells[0]]
                        )

                        docx.add_html(
                            row.cells[1],
                            study_endpoint.endpoint.name,
                            default_style="endpoint",
                        )
                        if len(row.cells[1].paragraphs) > 1:
                            # Remove first empty paragraph of cell
                            docx.delete_paragraph(row.cells[1].paragraphs[0])

                        if study_endpoint.timeframe:
                            docx.add_html(
                                row.cells[2],
                                study_endpoint.timeframe.name,
                                default_style="timeframe",
                            )
                            if len(row.cells[2].paragraphs) > 1:
                                # Remove first empty paragraph of cell
                                docx.delete_paragraph(row.cells[2].paragraphs[0])

                        units_text = self._endpoint_units_to_text(
                            study_endpoint.endpoint_units
                        )
                        docx.replace_content(row.cells[3], units_text, style="units")

        return docx

    def _endpoint_units_to_text(self, endpoint_units) -> str:
        separator = endpoint_units.separator

        if separator == ",":
            separator = ", "
        elif separator:
            separator = f" {separator} "
        else:
            separator = " "

        return separator.join(map(self._unit_name, endpoint_units.units))

    def _unit_name(self, unit_uid: str) -> str:
        unit: UnitDefinitionModel = self.units.get(unit_uid)
        if unit:
            return unit.name
        return unit_uid
