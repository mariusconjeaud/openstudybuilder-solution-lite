from typing import Any, Mapping

import yattag
from docx.shared import Inches
from pydantic import BaseModel, Field

from clinical_mdr_api.services.utils.docx_builder import DocxBuilder
from clinical_mdr_api.telemetry import trace_calls


class Ref(BaseModel):
    type: str | None = Field(..., title="Referenced item type")
    uid: str = Field(..., title="Referenced item uid")

    def __init__(self, type_=None, uid=None, **kwargs):
        if type_ is not None:
            kwargs["type"] = type_
        if uid is not None:
            kwargs["uid"] = uid
        super().__init__(**kwargs)


class TableCell(BaseModel):
    text: str = Field("", title="Text contents of cell")
    span: int = Field(1, title="Horizontal spanning of cell, 1 by default")
    style: str | None = Field(None, title="Associated style to cell")
    refs: list[Ref] | None = Field(None, title="Reference to item")
    footnotes: list[str] | None = Field(None, title="Referenced footnotes")
    vertical: bool | None = Field(None, title="Text text direction")

    def __init__(self, text=None, **kwargs):
        if text is not None:
            kwargs["text"] = text
        super().__init__(**kwargs)


class TableRow(BaseModel):
    cells: list[TableCell] = Field(default_factory=list, title="Table cells in the row")
    hide: bool = Field(False, title="Hide row from display")

    def __init__(self, cells=None, **kwargs):
        if cells is not None:
            kwargs["cells"] = cells
        super().__init__(**kwargs)


class SimpleFootnote(BaseModel):
    uid: str = Field(..., title="Unique id of footnote")  # TODO? Use Ref here?
    text_html: str = Field(..., title="HTML text of footnote")
    text_plain: str = Field(..., title="Plain text of footnote")


class TableWithFootnotes(BaseModel):
    rows: list[TableRow] = Field(default_factory=list, title="List of table rows")
    footnotes: dict[str, SimpleFootnote] | None = Field(
        None, title="Mapping of symbols and table footnotes"
    )
    num_header_rows: int = Field(0, title="Number of header rows")
    num_header_cols: int = Field(0, title="Number of header columns")
    title: str | None = Field(None, title="Table title (when rendered to HTML)")
    id: str | None = Field(None, title="Table id (when rendered to HTML)")


@trace_calls()
def table_to_docx(
    table: TableWithFootnotes, styles: Mapping[str, tuple[str, Any]] = None
) -> DocxBuilder:
    # assume horizontal table dimension from number of cells in first row
    num_cols = sum((c.span for c in table.rows[0].cells))

    # parses an empty template DOCX file into a helper class
    docx = DocxBuilder(
        styles=styles,
        landscape=True,
        margins=[0.5, 0.5, 0.5, 0.5],
    )

    # adds a table to the document
    x_table = docx.create_table(
        num_rows=sum(1 for row in table.rows if not row.hide),
        num_columns=num_cols,
    )

    # set width of first column
    x_table.columns[0].width = Inches(4)

    # cache porperty for performance issues, see github.com/python-openxml/python-docx/issues/174
    x_cells = x_table._cells
    x_rows = x_table.rows

    for r, t_row in enumerate((row for row in table.rows if not row.hide)):
        num_merge, merge_to = 0, None

        # set header row to repeat on each page
        if r < table.num_header_rows:
            docx.repeat_table_header(x_rows[r])

        for c, t_cell in enumerate(t_row.cells):
            x_cell = x_cells[r * num_cols + c]

            # merge with previous spanning cell
            if num_merge:
                merge_to.merge(x_cell)
                num_merge -= 1
                continue

            # skip invisible cells (should not get here if spans are coherent)
            if t_cell.span < 1:
                continue

            # when cell span > 1 merge the following N cells into this one
            num_merge = t_cell.span - 1
            merge_to = x_cell

            # all docx cells host one or more paragraph for contents
            x_para = x_cell.paragraphs[0]

            # set cell text in the paragraph
            if t_cell.text:
                x_para.text = t_cell.text

            # resolve style name and apply to paragraph
            style_name = styles.get(t_cell.style, [None])[0] if styles else None
            if style_name:
                x_para.style = style_name

            # set vertical text direction
            if t_cell.vertical:
                docx.set_vertical_cell_direction(x_cell, "btLr")

            # add a new run (like <span>) within the paragraph for each footnote symbol, with spacing in the run
            for symbol in t_cell.footnotes or []:
                run = x_para.add_run(f" {symbol}")
                run.font.superscript = True

    # add footnotes
    if table.footnotes:
        style_name = styles.get("footnote", [None])[0] if styles else None

        for symbol, footnote in table.footnotes.items():
            # each footnote is a new paragraph at the end of the document
            x_para = docx.document.add_paragraph(style=style_name)

            # footnote symbols into a run (like <span>) with superscript
            run = x_para.add_run(symbol)
            run.font.superscript = True

            # footnote text with glue and spacing into a distinct run
            x_para.add_run(f": {footnote.text_plain}")

    return docx


@trace_calls
def table_to_html(table: TableWithFootnotes) -> str:
    doc, tag, text, line = yattag.Doc().ttl()
    doc.asis("<!DOCTYPE html>")

    with tag("html", lang="en"):
        with tag("head"):
            if table.title:
                line("title", table.title)

        with tag("body"):
            attrs = {"id": table.id} if table.id else {}
            with tag("table", **attrs):
                with tag("thead"):
                    for row in table.rows[: table.num_header_rows]:
                        if row.hide:
                            continue

                        with tag("tr"):
                            for cell in row.cells:
                                if cell.span == 0:
                                    continue

                                with tag("th", **_cell_to_attrs(cell)):
                                    text(cell.text)
                                    for symbol in cell.footnotes or []:
                                        doc.asis("&nbsp;")
                                        line("sup", symbol)

                with tag("tbody"):
                    for row in table.rows[table.num_header_rows :]:
                        if row.hide:
                            continue

                        with tag("tr"):
                            for i, cell in enumerate(row.cells):
                                if cell.span == 0:
                                    continue

                                with tag(
                                    ("th" if i < table.num_header_cols else "td"),
                                    **_cell_to_attrs(cell),
                                ):
                                    text(cell.text)
                                    for symbol in cell.footnotes or []:
                                        doc.asis("&nbsp;")
                                        line("sup", symbol)

            if table.footnotes:
                with tag("dl", klass="footnotes"):
                    for symbol, footnote in table.footnotes.items():
                        line("dt", symbol)
                        with tag("dd"):
                            doc.asis(footnote.text_plain)

    return yattag.indent(doc.getvalue())


def _cell_to_attrs(cell):
    attrs = {}

    if cell.style:
        attrs["klass"] = cell.style

    if cell.span > 1:
        attrs["colspan"] = cell.span

    return attrs
