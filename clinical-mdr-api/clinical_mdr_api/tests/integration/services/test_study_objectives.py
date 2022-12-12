import logging
from io import BytesIO

import pytest
from bs4 import BeautifulSoup
from docx import Document

from clinical_mdr_api.tests.integration.utils.api import inject_and_clear_db
from clinical_mdr_api.tests.integration.utils.data_library import inject_base_data
from clinical_mdr_api.tests.integration.utils.method_library import generate_study_root
from clinical_mdr_api.tests.utils.checks import (
    assert_response_content_type,
    assert_response_status_code,
)

TEST_DB_NAME = __name__.rsplit(".", maxsplit=1)[-1].replace("_", "-")

log = logging.getLogger(__name__)


@pytest.fixture(scope="module")
def test_database():
    log.info(
        "test_database fixture: doing ugly magic to use database: %s", TEST_DB_NAME
    )
    inject_and_clear_db(TEST_DB_NAME)
    log.info(
        "test_database fixture: injecting base data into database: %s", TEST_DB_NAME
    )
    study = inject_base_data()
    return study


# pylint: disable=unused-argument,redefined-outer-name
def test_docx_response(app_client, test_database):
    study = generate_study_root()
    response = app_client.get(
        f"/studies/{study.uid}/study-objectives.docx", stream=True
    )
    assert_response_status_code(response, 200)
    assert_response_content_type(
        response,
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    doc = Document(BytesIO(response.raw.read()))
    assert len(doc.tables) == 1, "DOCX document must have exactly one table"
    table = doc.tables[0]
    assert len(table.columns) == 4, "expected 4 columns of table"
    assert len(table.rows) >= 1, "expected at least 1 row of table"


# pylint: disable=unused-argument,redefined-outer-name
def test_html_response(app_client, test_database):
    study = generate_study_root()
    response = app_client.get(
        f"/studies/{study.uid}/study-objectives.html", stream=True
    )
    assert_response_status_code(response, 200)
    assert_response_content_type(response, "text/html")
    doc = BeautifulSoup(response.raw, features="lxml")
    table = doc.find("table")
    assert table, "TABLE tag not found in document"
    assert table.get("id") == "ObjectivesEndpointsTable", "TABLE id mismatch"
    assert len(table.findAll("tr")), "TABLE has no TRs"
