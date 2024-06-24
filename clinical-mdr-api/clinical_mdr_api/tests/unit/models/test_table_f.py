# pylint: disable=no-member
from typing import Mapping

import bs4
import docx
import pytest
from docx.enum.style import WD_STYLE_TYPE

from clinical_mdr_api.services.utils.table_f import (
    SimpleFootnote,
    TableCell,
    TableRow,
    TableWithFootnotes,
    table_to_docx,
    table_to_html,
)

DOCX_TEXT_DIRECTION_VALUE = (
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
)
DOCX_TEXT_DIRECTIION_TAG = (
    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}textDirection"
)

DOCX_STYLES = {
    "table": ("My Table Style", WD_STYLE_TYPE.TABLE),
    "head": ("Head Style", WD_STYLE_TYPE.PARAGRAPH),
    "data": ("Data Style", WD_STYLE_TYPE.PARAGRAPH),
}

TEST_TABLE = TableWithFootnotes(
    rows=[
        TableRow(
            hide=False,
            cells=[
                TableCell(text="hi", style="hi hi", vertical=True, footnotes=["a"]),
                TableCell(text="Hello", span="2", style="hi"),
                TableCell(style="foo", span=0),
                TableCell(text="Hi World"),
            ],
        ),
        TableRow(
            hide=True,
            cells=[
                TableCell("Blah Text", style="head2", footnotes=["x", "y"]),
                TableCell("more head", span=2),
                TableCell(span=0),
                TableCell("Also column"),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("some Text", style="head", footnotes=["hello", "z13"]),
                TableCell("More Head", span=2, footnotes=["z"]),
                TableCell(span=0),
                TableCell("also a head"),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("This goes TH", style="head"),
                TableCell("More TH", style="head2"),
                TableCell("data #1 cell"),
                TableCell("data cell #2", style="data"),
            ],
        ),
        TableRow(
            hide=True,
            cells=[
                TableCell("some hidden"),
                TableCell(span="2"),
                TableCell(span=0),
                TableCell("contents blah", footnotes=["hello", "a"]),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("more Data th", vertical=True),
                TableCell(span=2),
                TableCell(span=0),
                TableCell("foo", footnotes=["z13", "hello"]),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("foo-Bar", span=2, footnotes=["c", "X"]),
                TableCell(span=0),
                TableCell(),
                TableCell("X", footnotes=["z13", "hello"]),
            ],
        ),
        TableRow(
            hide=False,
            cells=[
                TableCell("even more data"),
                TableCell("second th row"),
                TableCell("X", style="data", footnotes=["a"]),
                TableCell("data"),
            ],
        ),
    ],
    num_header_rows=3,
    num_header_cols=2,
    title="Hello Table",
    footnotes={
        "a": SimpleFootnote(uid="fn01", text_html="not used", text_plain="Footote ej"),
        "hello": SimpleFootnote(
            uid="foot76",
            text_plain="Hello footnotes here and there",
            text_html="not displayed",
        ),
        "z13": SimpleFootnote(
            uid="foot42",
            text_plain="More footnotes after numbering",
            text_html="not shown",
        ),
    },
)


EXPECTED_HTML = """<!DOCTYPE html>
<html lang="en">
  <head>
    <title>Hello Table</title>
  </head>
  <body>
    <table>
      <thead>
        <tr>
          <th class="hi hi">hi&nbsp;<sup>a</sup></th>
          <th class="hi" colspan="2">Hello</th>
          <th>Hi World</th>
        </tr>
        <tr>
          <th class="head">some Text&nbsp;<sup>hello</sup>&nbsp;<sup>z13</sup></th>
          <th colspan="2">More Head&nbsp;<sup>z</sup></th>
          <th>also a head</th>
        </tr>
      </thead>
      <tbody>
        <tr>
          <th class="head">This goes TH</th>
          <th class="head2">More TH</th>
          <td>data #1 cell</td>
          <td class="data">data cell #2</td>
        </tr>
        <tr>
          <th>more Data th</th>
          <th colspan="2"></th>
          <td>foo&nbsp;<sup>z13</sup>&nbsp;<sup>hello</sup></td>
        </tr>
        <tr>
          <th colspan="2">foo-Bar&nbsp;<sup>c</sup>&nbsp;<sup>X</sup></th>
          <td></td>
          <td>X&nbsp;<sup>z13</sup>&nbsp;<sup>hello</sup></td>
        </tr>
        <tr>
          <th>even more data</th>
          <th>second th row</th>
          <td class="data">X&nbsp;<sup>a</sup></td>
          <td>data</td>
        </tr>
      </tbody>
    </table>
    <dl class="footnotes">
      <dt>a</dt>
      <dd>Footote ej</dd>
      <dt>hello</dt>
      <dd>Hello footnotes here and there</dd>
      <dt>z13</dt>
      <dd>More footnotes after numbering</dd>
    </dl>
  </body>
</html>"""


def test_table_to_html_expected():
    html = table_to_html(TEST_TABLE)
    assert html == EXPECTED_HTML


@pytest.mark.parametrize("test_table", [TEST_TABLE])
def test_table_to_html(test_table: TableWithFootnotes):
    """Tests table_to_html by comparing table contents to TableWithFootnotes input"""

    html = table_to_html(TEST_TABLE)

    # THEN body is a parseable HTML document
    # html.parser does not supplement the document with additional elements for the sake of well-formedness
    doc = bs4.BeautifulSoup(html, features="html.parser")

    # THEN document has HTML tag
    assert doc.find("html"), "<HTML> tag not found"

    # THEN the document has HTML > HEAD > TITLE tags
    assert (title := doc.select_one("html > head > title")), "<TITLE> tag not found"
    # THEN document title matches
    assert title.text == test_table.title, "title mismatch"

    # THEN the document has TABLE
    assert (table := doc.select_one("html > body > table")), "<TABLE> tag not found"

    compare_html_table(table, test_table)

    if test_table.footnotes:
        compare_html_footnotes(doc, test_table)


def compare_html_table(table: bs4.element.Tag, test_table: TableWithFootnotes):
    """Compares HTML table with TableWithFootnotes by cell contents and properties"""

    # THEN number of table rows match
    expected_num_rows = sum(1 for row in test_table.rows if not row.hide)
    assert (
        len(table.find_all("tr")) == expected_num_rows
    ), "unexpected number of rows<tr>"

    # THEN number of header rows match
    num_header_rows = sum(
        1 for row in test_table.rows[: test_table.num_header_rows] if not row.hide
    )
    assert (
        len(table.select("thead > tr")) == num_header_rows
    ), "unexpected number of header rows <thead><tr>"

    row: TableRow
    # pylint: disable=invalid-name
    # iterate over index, <tr>, TableRow
    for r, (tr, row) in enumerate(
        zip(table.find_all("tr"), (row for row in test_table.rows if not row.hide))
    ):
        if r < num_header_rows:
            assert (
                tr.parent.name.lower() == "thead"
            ), f"expected parent of row {r} is <thead>"
        else:
            assert (
                tr.parent.name.lower() == "tbody"
            ), f"expected parent of row {r} is <tbody>"

        cells = [cell for cell in row.cells if cell.span]

        # THEN number of cells in rows match
        assert len(tds := tr.find_all(["th", "td"])) == len(
            cells
        ), "unexpected number of cells in row i"

        cell: TableCell
        span = 0
        # pylint: disable=invalid-name
        # iterate over index, <td/th>, TableCell
        for c, (td, cell) in enumerate(zip(tds, cells)):
            if r < num_header_rows or span < test_table.num_header_cols:
                assert (
                    td.name.lower() == "th"
                ), f"expected element in row {r} cell {c} is <th>"
            else:
                assert (
                    td.name.lower() == "td"
                ), f"expected element in row {r} cell {c} is <th>"
            span += cell.span

            next_string = next(td.stripped_strings, "")
            # THEN cell text matches
            assert (
                next_string == cell.text
            ), f"cell text does not match in row {r} column {c}"

            # THEN horizontal cell spanning matches
            assert td.get("colspan", "1") == str(
                cell.span
            ), f"cell span does not match in row {r} column {c}"

            # THEN cell styling matches
            style = " ".join(td.get("class", [])) or None
            assert (
                style == cell.style
            ), f"cell style does not match in row {r} column {c}"

            if cell.footnotes:
                # THEN footnote symbols found in cell as superscript
                assert (
                    sups := td.find_all("sup")
                ), f"cell missing <sup> for footnotes in row {r} column {c}"

                sup_texts = [sup.text.strip() for sup in sups]
                assert (
                    sup_texts == cell.footnotes
                ), f"unexpected superscript text of footnote symbols in row {r} column {c}"


# pylint: disable=invalid-name
def compare_html_footnotes(doc: bs4.BeautifulSoup, test_table: TableWithFootnotes):
    """Compares footnote <dl> with footnotes of TableWithFootnotes"""

    assert (
        dl := doc.find("dl", class_="footnotes")
    ), "no <DL> tag found with .footnotes class"

    assert len(dts := dl.find_all("dt")) == len(
        dds := dl.find_all("dd")
    ), "number of <DT> and <DD> tags mismatch"

    # THEN number of footnotes match with footnotes list
    assert len(dts) == len(test_table.footnotes), "mumber of footnotes mismatch"

    for r, (dt, dd, symbol) in enumerate(zip(dts, dds, test_table.footnotes.keys())):
        footnote = test_table.footnotes[symbol]

        # THEN footnote symbols match with footnotes list
        assert dt.text == symbol, f"<dt> text does not match symbol in footnote row {r}"
        # Then footnote text matches with footnote list
        assert (
            dd.text == footnote.text_plain
        ), f"<dd> text does not match plain text in footnote row {r}"


@pytest.mark.parametrize("test_table", [TEST_TABLE])
def test_table_to_docx(test_table: TableWithFootnotes):
    """Tests table_to_docx() by comparing DOCX document to TableWithFootnotes input"""

    docx_doc: docx.Document = table_to_docx(test_table, styles=DOCX_STYLES).document

    # THEN the document contains exactly one table
    assert len(docx_doc.tables) == 1, "expected exactly 1 table in DOCX SoA"

    compare_docx_table(docx_doc.tables[0], test_table, DOCX_STYLES)

    if test_table.footnotes:
        compare_docx_footnotes(docx_doc, test_table.footnotes, DOCX_STYLES)


def compare_docx_table(
    tablex: docx.table.Table,
    test_table: TableWithFootnotes,
    docx_styles: Mapping[str, tuple[str, any]],
):
    """Compares DOCX table with TableWithFootnotes by column contents and properties"""

    # THEN number of table rows match
    expected_num_rows = sum(1 for row in test_table.rows if not row.hide)
    assert len(tablex.rows) == expected_num_rows, "unexpected number of table rows"

    # THEN number of table columns match
    expected_num_cols = sum(cell.span for cell in test_table.rows[-1].cells)
    assert (
        len(tablex.columns) == expected_num_cols
    ), "unexpected number of table columns"

    row: TableRow
    rowx: docx.table._Row
    for row_idx, (rowx, row) in enumerate(
        zip(tablex.rows, (row for row in test_table.rows if not row.hide))
    ):
        cell: TableCell
        cellx: docx.table._Cell
        for col_idx, (cellx, cell) in enumerate(zip(rowx.cells, row.cells)):
            if not cell.span:
                # In DocX, despite merging cells, the number of cells in a row is constant across the table, and the
                # contents of the first merged member are repeated for all members.
                # So when we encounter a span=0 cell in JSON, it's a subsequent member of a merged cell,
                # we skip checking the contents and formatting of the corresponding cell in DocX.
                continue

            # first paragraph of cell
            parax0: docx.text.paragraph.Paragraph = cellx.paragraphs[0]

            # THEN cell text matches
            textx = parax0.runs[0].text if parax0.runs else parax0.text
            assert (
                textx == cell.text
            ), f"cell text does not match in row {row_idx} column {col_idx}"

            if cell.vertical:
                # THEN some header cell texts are oriented vertical
                tcpr = next(
                    (
                        tcPr
                        for tcPr in cellx._tc.tcPr
                        if tcPr.tag == DOCX_TEXT_DIRECTIION_TAG
                    ),
                    None,
                )
                assert (
                    tcpr is not None
                ), f"vertical cell direction not set on row {row_idx} column {col_idx}"
                assert (
                    tcpr.get(DOCX_TEXT_DIRECTION_VALUE) == "btLr"
                ), f"vertical cell direction 'btLr' expected row {row_idx} column {col_idx}"

            if cell.style in docx_styles:
                # THEN cell styling matches
                expected_style_name = docx_styles.get(cell.style, ["Normal"])[0]
                assert (
                    parax0.style.name == expected_style_name
                ), f"cell style does not match in row {row_idx} column {col_idx}"

            if cell.footnotes:
                # THEN number of footnote symbols appended to cell text matches
                assert len(parax0.runs[1:]) == len(
                    cell.footnotes
                ), f"runber of runs does not match number of footnotes in row {row_idx} column {col_idx}"

                # THEN footnote symbols match
                assert [
                    runx.text.strip() for runx in parax0.runs[1:]
                ] == cell.footnotes, f"unexpected superscript text of footnote symbols in row {row_idx} column {col_idx}"

                # THEN footnote symbols are in superscript
                assert all(
                    runx.font.superscript for runx in parax0.runs[1:]
                ), f"not all footnote symbols are superscript in row {row_idx} column {col_idx}"


def compare_docx_footnotes(
    docx_doc: docx.Document,
    footnotes: dict[str, SimpleFootnote],
    docx_styles: Mapping[str, tuple[str, any]],
):
    """Compares DOCX paragraphs with footnotes of TableWithFootnotes"""

    # Table is not a paragraph, so all top-level paragraps of a SoA DOCX document are part of footnote listing
    footnote_paras = [parax for parax in docx_doc.paragraphs if parax.text.strip()]

    # THEN document lists footnotes
    assert len(footnote_paras), "footnotes listing was not found in DOCX document"

    # THEN number of footnotes match
    assert len(footnote_paras) == len(footnotes), "mumber of footnotes mismatch"

    for row_idx, (parax, (symbol, footnote)) in enumerate(
        zip(footnote_paras, footnotes.items())
    ):
        # THEN footnote paragraph has styling
        assert parax.style.name == docx_styles.get("footnote", ["Normal"])[0]

        # THEN footnote symbols match with footnotes list
        textx = parax.runs[0].text
        assert textx == symbol, f"footnote symbol does not match in row {row_idx}"

        # THEN footnote symbols are in superscript
        assert parax.runs[
            0
        ].font.superscript, f"footnote symbol is not superscript in row {row_idx}"

        # Then footnote text matches with footnotes list
        textx = parax.runs[1].text.lstrip(": ")
        assert (
            textx == footnote.text_plain
        ), f"footnote text does not match in row {row_idx}"
